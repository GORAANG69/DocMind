"""DOCX parser using python-docx."""
from pathlib import Path

from docx import Document as DocxDocument

from parsers.base_parser import BaseParser


class DocxParser(BaseParser):
    """Extract text from DOCX files."""

    def extract_text(self, file_path: Path) -> str:
        try:
            doc = DocxDocument(str(file_path))
        except Exception as exc:
            raise ValueError(f"Cannot open DOCX '{file_path.name}': {exc}") from exc

        parts: list[str] = []

        # Paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        if not parts:
            raise ValueError(f"DOCX '{file_path.name}' contains no extractable text.")

        return "\n".join(parts)

    @staticmethod
    def supported_extensions() -> list[str]:
        return [".docx"]
